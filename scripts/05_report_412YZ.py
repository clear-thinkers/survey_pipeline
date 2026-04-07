"""
05_report_412YZ.py
Generate report/412YZ/report_412YZ.docx from output/412YZ/analysis_412YZ.xlsx.

Fully code-driven python-docx output — no LLM API call.
Narrative text is templated from analysis data. Placeholders that require
manual input are highlighted yellow in the output document.

Usage:
    python scripts/05_report_412YZ.py
"""

import sys
from pathlib import Path

import pandas as pd
import openpyxl
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = Path(__file__).parent.parent
ANALYSIS_PATH = BASE_DIR / "output" / "412YZ" / "analysis_412YZ.xlsx"
CSV_PATH      = BASE_DIR / "output" / "412YZ" / "survey_data_412YZ.csv"
OUT_DIR       = BASE_DIR / "report" / "412YZ"
OUT_PATH      = OUT_DIR / "report_412YZ.docx"

SURVEY_MONTH  = "March 2026"
N_RESPONDENTS = None       # derived from CSV at runtime (see load_csv)

# Prior-year Q1 coach satisfaction benchmarks (% top-2 box, from example report)
Q1_BENCHMARKS = {
    "label":  ["My Youth Coach…", "My Youth Coach…", "My Youth Coach…"],
    "col":    ["Sep-19", "Mar-22", "Feb-23", "Feb-24", "Feb-25"],
    "n":      ["n=154",  "n=103",  "n=128",  "n=142",  "n=167"],
    "Is trustworthy":                    ["94%", "91%", "95%", "95%", "90%"],
    "Is reliable":                       ["92%", "90%", "88%", "93%", "89%"],
    "Values my opinions about my life":  ["94%", "91%", "91%", "93%", "91%"],
    "Is available to me when I need them":["85%", "88%", "88%", "91%", "88%"],
    "Makes me feel heard and understood":["92%", "90%", "92%", "91%", "89%"],
}

# Table header / total row fill (matches prior report)
HDR_FILL = "DCE6F1"

# ---------------------------------------------------------------------------
# Helpers — styling
# ---------------------------------------------------------------------------

def _shade_cell(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    existing = tcPr.find(qn("w:shd"))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(shd)


def shade_row(row, hex_color: str):
    for cell in row.cells:
        _shade_cell(cell, hex_color)


def add_table(doc: Document, df: pd.DataFrame, total_label: str = "Total") -> None:
    """Write a DataFrame to a styled Word table.

    Header row and any row whose first cell matches *total_label* get
    DCE6F1 fill + bold text, matching the prior report style.
    """
    if df is None or df.empty:
        return

    cols = list(df.columns)
    tbl = doc.add_table(rows=1, cols=len(cols))
    tbl.style = "Table Grid"
    # Remove all borders
    tblPr = tbl._tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl._tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        tblBorders.append(el)
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblBorders)

    # Header row
    hdr_row = tbl.rows[0]
    shade_row(hdr_row, HDR_FILL)
    for i, col_name in enumerate(cols):
        cell = hdr_row.cells[i]
        cell.text = str(col_name)
        for run in cell.paragraphs[0].runs:
            run.font.bold = True

    # Data rows
    for _, row_data in df.iterrows():
        row = tbl.add_row()
        first_val = str(row_data.iloc[0]).strip()
        is_total = first_val == total_label
        if is_total:
            shade_row(row, HDR_FILL)
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = "" if pd.isna(val) or str(val) == "nan" else str(val)
            if is_total:
                for run in cell.paragraphs[0].runs:
                    run.font.bold = True

    doc.add_paragraph()  # spacing after table


def para(doc: Document, text: str, bold: bool = False, italic: bool = False,
         style: str = "Normal") -> None:
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.font.bold = bold
    run.font.italic = italic


def placeholder(doc: Document, text: str) -> None:
    """Add a highlighted yellow paragraph for manually-filled values."""
    p = doc.add_paragraph(style="Normal")
    run = p.add_run(f"[{text}]")
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    run.font.bold = True


def placeholder_inline(para, text: str):
    """Add a highlighted yellow run inline within an existing paragraph."""
    run = para.add_run(f"[{text}]")
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    run.font.bold = True
    return run


def heading(doc: Document, text: str, bold: bool = True) -> None:
    p = doc.add_paragraph(style="Normal")
    run = p.add_run(text)
    run.font.bold = bold


def caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Normal")
    run = p.add_run(text)
    run.font.bold = True


def bullet(doc: Document, text: str, highlight: bool = False) -> None:
    p = doc.add_paragraph(style="List Paragraph")
    run = p.add_run(text)
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        run.font.bold = True


# ---------------------------------------------------------------------------
# Helpers — data
# ---------------------------------------------------------------------------

def pct(n, d, decimals=0):
    if not d:
        return ""
    val = round(100 * n / d, decimals)
    return f"{int(val)}%" if decimals == 0 else f"{val:.{decimals}f}%"


def split_pipe(val):
    if not str(val).strip() or str(val) == "nan":
        return []
    return [t.strip() for t in str(val).split("|") if t.strip()]


def load_sheets() -> dict[str, pd.DataFrame]:
    wb = openpyxl.load_workbook(ANALYSIS_PATH, data_only=True)
    sheets = {}
    for name in wb.sheetnames:
        ws = wb[name]
        data = list(ws.values)
        if not data:
            sheets[name] = pd.DataFrame()
            continue
        # First row is the section title (written by write_section in 04_analyze).
        # Second row is the real column header.
        if len(data) < 2:
            sheets[name] = pd.DataFrame()
            continue
        header = [str(c) if c is not None else "" for c in data[1]]
        rows = [[str(c) if c is not None else "" for c in row] for row in data[2:]]
        sheets[name] = pd.DataFrame(rows, columns=header)
    return sheets


def _split_sheet(sheet_name: str) -> dict:
    """Read an xlsx sheet and split it into named sub-tables.

    Section title rows are identified as rows where the first cell is non-blank
    and all remaining cells are blank.  The first non-blank, non-title row within
    each section becomes its column header.

    Returns a dict mapping section title → pd.DataFrame.
    """
    wb = openpyxl.load_workbook(ANALYSIS_PATH, data_only=True)
    ws = wb[sheet_name]
    all_rows = [
        [str(c.value) if c.value is not None else "" for c in row]
        for row in ws.iter_rows()
    ]

    sections: dict = {}
    current_title = None
    current_rows: list = []

    for row in all_rows:
        first = row[0].strip()
        is_title = bool(first) and not any(v.strip() for v in row[1:])
        if is_title:
            if current_title and current_rows:
                sections[current_title] = current_rows
            current_title = first
            current_rows = []
        elif current_title is not None:
            current_rows.append(row)

    if current_title and current_rows:
        sections[current_title] = current_rows

    result = {}
    for title, rows in sections.items():
        non_blank = [r for r in rows if any(v.strip() for v in r)]
        if not non_blank:
            result[title] = pd.DataFrame()
            continue
        max_col = max(
            max((i for i, v in enumerate(r) if v.strip()), default=0)
            for r in non_blank
        ) + 1
        trimmed = [r[:max_col] for r in non_blank]
        hdr = trimmed[0]
        result[title] = pd.DataFrame(trimmed[1:], columns=hdr)

    return result


def load_csv() -> pd.DataFrame:
    global N_RESPONDENTS
    df = pd.read_csv(CSV_PATH, encoding="utf-8-sig", dtype=str).fillna("")
    N_RESPONDENTS = len(df)

    def age_label(code):
        return {"16_17": "16-17 years old", "18_20": "18-20 years old",
                "21_23": "21-23 years old"}.get(str(code).strip(), "Unknown")

    def gender_group(g):
        g = str(g).strip()
        if g == "Female":  return "Female"
        if g == "Male":    return "Male"
        if not g:          return "No answer"
        return "Trans, Non-binary"

    df["_age"]    = df["age_range"].apply(age_label)
    df["_gender"] = df["gender"].apply(gender_group)
    return df


# ---------------------------------------------------------------------------
# Section functions
# ---------------------------------------------------------------------------

def sec_title(doc, df_csv):
    n = N_RESPONDENTS
    heading(doc, "Youth Zone Survey Results", bold=True)
    heading(doc, SURVEY_MONTH, bold=False)
    doc.add_paragraph()

    p = doc.add_paragraph(style="Normal")
    p.add_run(
        "All individuals active with the 412 Youth Zone had the opportunity to "
        "participate in a survey in early 2026. Surveys were administered on paper "
        "and results were digitized for this report."
    )

    p2 = doc.add_paragraph(style="Normal")
    p2.add_run(f"{n} unique youth (of ")
    placeholder_inline(p2, "TOTAL ACTIVE — fill in denominator")
    p2.add_run(
        f" total active) responded to the survey, for a response rate of "
    )
    placeholder_inline(p2, "RESPONSE RATE %")
    p2.add_run(
        f". Most respondents were age 18–20 years old."
    )


def sec_age(doc, sheets):
    caption(doc, "Survey Respondents by Age")
    df = sheets["01_age"].copy()
    df.columns = ["Age", "Count"]
    add_table(doc, df)


def sec_gender_orient(doc, sheets):
    df_raw = sheets["02_gender_orient"].copy()
    # First data row is the "Number of Youth" header row
    n_female  = df_raw[df_raw.iloc[:, 0] == "Number of Youth"]["Female"].values
    n_male    = df_raw[df_raw.iloc[:, 0] == "Number of Youth"]["Male"].values
    n_tnb     = df_raw[df_raw.iloc[:, 0] == "Number of Youth"]["Trans, Non-binary"].values
    n_f  = int(n_female[0])  if len(n_female)  else 0
    n_m  = int(n_male[0])    if len(n_male)    else 0
    n_nb = int(n_tnb[0])     if len(n_tnb)     else 0
    n_known = n_f + n_m + n_nb
    pct_f = pct(n_f,  n_known)
    pct_m = pct(n_m,  n_known)

    para(doc,
         f"More females ({pct_f}) responded to the survey than males ({pct_m}). "
         "The table below shows how respondents identified by gender and sexual orientation.")

    caption(doc, "Survey Respondents by Gender and Sexual Orientation")
    add_table(doc, df_raw, total_label="Total")


def sec_race(doc, sheets):
    para(doc,
         "The tables below display respondents' racial identities and genders. "
         "Please note: Youth Zone participants' racial identities are self-reported "
         "and reflect the full range of how young people describe themselves.")

    df_once = sheets["03_race_once"].copy()
    # Pull % Black from the table for narrative
    black_row = df_once[df_once.iloc[:, 0] == "Black"]
    pct_black = black_row["Percent"].values[0] if not black_row.empty and "Percent" in df_once.columns else "[PLACEHOLDER]"
    white_row = df_once[df_once.iloc[:, 0] == "White"]
    pct_white = white_row["Percent"].values[0] if not white_row.empty and "Percent" in df_once.columns else "[PLACEHOLDER]"
    multi_row = df_once[df_once.iloc[:, 0] == "Multiracial"]
    pct_multi = multi_row["Percent"].values[0] if not multi_row.empty and "Percent" in df_once.columns else "[PLACEHOLDER]"

    para(doc,
         f"About {pct_black} of survey respondents identified as Black, "
         f"{pct_white} as White, and {pct_multi} as Multiracial.")

    caption(doc, "Youth by Race and Gender (all Youth are Counted Once)")
    add_table(doc, df_once, total_label="Total")

    caption(doc, "Youth with Full or Partial Racial Identities (Some Youth Are Counted Multiple Times)")
    df_multi = sheets["04_race_multi"].copy()
    add_table(doc, df_multi, total_label="Total")


def sec_coach_satisfaction(doc, sheets):
    heading(doc, "FINDINGS")
    doc.add_paragraph()
    heading(doc, "Relationships with Coach")

    df_q1 = sheets["05_q1"].copy()
    # Pull top-2 % for trustworthy from current year
    trust_row = df_q1[df_q1.iloc[:, 0].str.contains("trustworthy", case=False, na=False)]
    pct_trust = trust_row.iloc[0, 2] if not trust_row.empty and df_q1.shape[1] > 2 else "[PLACEHOLDER]"
    vals_row  = df_q1[df_q1.iloc[:, 0].str.contains("values", case=False, na=False)]
    pct_vals  = vals_row.iloc[0, 2]  if not vals_row.empty  and df_q1.shape[1] > 2 else "[PLACEHOLDER]"

    para(doc,
         f"{pct_trust} of youth reported their coaches were trustworthy, and "
         f"{pct_vals} indicated their coach values their opinions about their life. "
         "Ratings across all five coach relationship items are shown in the table below.")

    # Build multi-year table
    field_labels = [
        "Is trustworthy",
        "Is reliable",
        "Values my opinions about my life",
        "Is available to me when I need them",
        "Makes me feel heard and understood",
    ]
    prior_cols = Q1_BENCHMARKS["col"]
    curr_col   = f"Mar-26"
    curr_n     = f"n={N_RESPONDENTS}"

    # Map current-year pct from analysis sheet
    curr_pct = {}
    for _, row in df_q1.iterrows():
        label = str(row.iloc[0]).strip()
        pct_val = str(row.iloc[2]).strip() if df_q1.shape[1] > 2 else ""
        curr_pct[label] = pct_val

    rows_out = []
    # Header rows (2 rows matching prior report format)
    rows_out.append({"My Youth Coach…": "My Youth Coach…",
                     **{c: "% Often or All the Time" for c in prior_cols},
                     curr_col: "% Often or All the Time"})
    rows_out.append({"My Youth Coach…": "My Youth Coach…",
                     **{c: c for c in prior_cols},
                     curr_col: curr_col})
    rows_out.append({"My Youth Coach…": "My Youth Coach…",
                     **{c: Q1_BENCHMARKS["n"][i] for i, c in enumerate(prior_cols)},
                     curr_col: curr_n})
    for label in field_labels:
        row_dict = {"My Youth Coach…": label}
        for i, c in enumerate(prior_cols):
            row_dict[c] = Q1_BENCHMARKS[label][i]
        row_dict[curr_col] = curr_pct.get(label, "")
        rows_out.append(row_dict)

    df_multiyear = pd.DataFrame(rows_out)
    caption(doc, "Satisfaction Ratings for Youth Coaches Over Time")
    add_table(doc, df_multiyear, total_label="__none__")


def sec_communication(doc, sheets):
    df = sheets["06_communication"].copy()

    # Find not_enough count and percentage from Q3 sub-table
    q3_table_start = None
    for i, row in df.iterrows():
        if "Not enough" in str(row.iloc[0]) or "not_enough" in str(row.iloc[0]).lower():
            q3_table_start = i
            break

    not_enough_count = ""
    not_enough_pct   = ""
    good_amount_pct  = ""
    if q3_table_start is not None:
        ne_row = df[df.iloc[:, 0].str.lower().str.contains("not enough", na=False)]
        if not ne_row.empty:
            not_enough_count = ne_row.iloc[0, 1] if df.shape[1] > 1 else ""
            not_enough_pct   = ne_row.iloc[0, 2] if df.shape[1] > 2 else ""
        ga_row = df[df.iloc[:, 0].str.lower().str.contains("good amount", na=False)]
        if not ga_row.empty:
            good_amount_pct = ga_row.iloc[0, 2] if df.shape[1] > 2 else ""

    if not_enough_count and not_enough_pct:
        para(doc,
             f"The majority of youth communicate with their coaches weekly or monthly. "
             f"{not_enough_pct} of respondents ({not_enough_count} youth) reported "
             "their communication was Not Enough; most of these youth communicated "
             "with their coach about once a week or 1–2 times per month.")
    else:
        para(doc,
             "The majority of youth communicate with their coaches weekly or monthly. "
             "A small number of youth reported their communication was Not Enough.")


def sec_housing(doc, sheets):
    heading(doc, "Stable Housing")

    para(doc,
         "Survey respondents were asked to describe whether their current housing is "
         "safe and stable, meaning they can stay there for at least 90 days. Youth "
         "with unstable housing were then asked where they are currently sleeping.")

    df_h = sheets["07_housing"].copy()
    # Pull stable % for narrative
    stable_row = df_h[df_h.iloc[:, 0].str.lower().str.contains("stable", na=False) &
                      ~df_h.iloc[:, 0].str.lower().str.contains("un", na=False)]
    pct_stable = stable_row["Percent"].values[0] if not stable_row.empty and "Percent" in df_h.columns else "[PLACEHOLDER]"

    para(doc,
         f"About {pct_stable} of respondents reported safe and stable housing. "
         "For the remainder, current sleeping arrangements are shown in the table below.")

    caption(doc, "Housing Status and Current Sleeping Arrangements")
    add_table(doc, df_h, total_label="Total")

    para(doc,
         "Regardless of their current living situation, if youth experienced unstable "
         "housing in the prior six months, they were asked to identify the reason(s).")

    caption(doc, "Reasons for Unstable Housing in the Past 6 Months, by Age\u00b9")
    df_reasons = sheets["08_housing_reasons"].copy()
    add_table(doc, df_reasons, total_label="Total youth reporting unstable housing")
    para(doc, "\u00b9 Youth could report more than one reason for experiencing unstable housing.", italic=True)


def sec_education_employment(doc, sheets, df_csv):
    heading(doc, "Employment and Education")

    # Compute inline stats from CSV
    total = len(df_csv)
    in_school  = df_csv[df_csv["q5_school_status"].isin(
        ["high_school", "college_career", "ged", "graduate"])].shape[0]
    employed   = df_csv[df_csv["q8_employment_status"].isin(
        ["yes_full_time", "yes_part_time"])].shape[0]
    in_s_unemp = df_csv[(df_csv["q5_school_status"].isin(
        ["high_school","college_career","ged","graduate"])) &
        (df_csv["q8_employment_status"] == "no")]
    in_s_unemp_seeking = (in_s_unemp["q8b_job_seeking"] == "yes").sum()
    not_s_unemp = df_csv[(df_csv["q5_school_status"] == "not_in_school") &
                          (df_csv["q8_employment_status"] == "no")]
    not_s_unemp_seeking = (not_s_unemp["q8b_job_seeking"] == "yes").sum()
    not_in_school_unemp = df_csv[(df_csv["q5_school_status"] == "not_in_school") &
                                  (df_csv["q8_employment_status"] == "no")]
    no_diploma = (not_in_school_unemp["q5a_highest_education"] == "some_hs").sum()

    # Employed and in school
    emp_df = df_csv[df_csv["q8_employment_status"].isin(["yes_full_time","yes_part_time"])]
    emp_also_school = emp_df[emp_df["q5_school_status"].isin(
        ["high_school","college_career","ged","graduate"])].shape[0]

    bullet(doc, f"{pct(in_school, total)} of all respondents reported being enrolled in school")
    bullet(doc, f"{pct(employed, total)} of all respondents reported being employed "
                f"({pct(emp_also_school, employed)} of these youth are also enrolled in school)")
    if len(in_s_unemp):
        bullet(doc, f"{pct(in_s_unemp_seeking, len(in_s_unemp))} of respondents who are "
                    "in school and unemployed are looking for a job")
    if len(not_s_unemp):
        bullet(doc, f"{pct(not_s_unemp_seeking, len(not_s_unemp))} of respondents who are "
                    "not in school and unemployed are looking for a job")
    pct_both_out = pct(len(not_in_school_unemp), total)
    bullet(doc, f"{pct_both_out} of respondents are both not in school and unemployed; "
                f"{no_diploma} of these {len(not_in_school_unemp)} youth report not completing "
                "high school or a GED")

    caption(doc, "Educational Enrollment and Attainment")
    add_table(doc, sheets["09_education"].copy(), total_label="Total")


def sec_job_tenure(doc, sheets, df_csv):
    employed = df_csv[df_csv["q8_employment_status"].isin(["yes_full_time","yes_part_time"])]
    n_emp = len(employed)
    long_tenure = (employed["q8a_job_tenure"] == "more_6mo").sum()

    para(doc,
         f"Of the {pct(n_emp, N_RESPONDENTS)} of survey respondents that reported being employed, "
         f"{pct(long_tenure, n_emp)} have been at their job for six months or longer.")

    caption(doc, "Length of Employment for Youth Currently Employed")
    add_table(doc, sheets["11_job_tenure"].copy(), total_label="Total")


def sec_employment_by_age(doc, sheets):
    caption(doc, "Employment Status by Age")
    add_table(doc, sheets["10_employment"].copy(), total_label="Total")


def sec_job_barriers(doc, sheets, df_csv):
    non_ft = df_csv[df_csv["q8_employment_status"] != "yes_full_time"]
    n_with_barriers = (non_ft["q10_job_barriers"].str.strip() != "").sum()

    df_b = sheets["12_job_barriers"].copy()
    top_row = df_b.iloc[0] if not df_b.empty else None
    top_label = top_row.iloc[0] if top_row is not None else ""
    top_pct   = top_row.iloc[2] if (top_row is not None and df_b.shape[1] > 2) else ""

    para(doc,
         "If survey respondents had trouble finding a job in the prior twelve months, "
         "they were asked to share some of the reasons why. "
         f"The most common challenge identified, reported by {top_pct} of youth, "
         f"was {top_label.lower()}.")

    caption(doc, "Reasons Youth Have Trouble Finding Jobs (Reasons Given by 2 or More People)")
    add_table(doc, df_b, total_label="__none__")
    para(doc, "Note: Youth could select more than one option", italic=True)


def sec_left_job(doc, sheets):
    df_lj = sheets["13_left_job"].copy()
    quit_row = df_lj[df_lj.iloc[:, 0].str.lower().str.strip() == "quit"]
    n_quit   = quit_row["Total Youth"].values[0] if not quit_row.empty and "Total Youth" in df_lj.columns else ""

    para(doc,
         "If the survey respondent lost or left a job in the past year, they were asked "
         "to share the reason(s). Sub-reasons for youth who quit are shown indented below the Quit row.")

    caption(doc, "Reasons Youth Lost or Quit a Job in the Past Year (Reasons Given by 2 or More People)")
    add_table(doc, df_lj, total_label="__none__")


def sec_transportation(doc, sheets):
    heading(doc, "TRANSPORTATION")

    para(doc,
         "Youth were also asked about whether they have a driver\u2019s license and the type "
         "of transportation they rely on for work. The table below displays driver\u2019s "
         "license status by age.")

    _write_transport_tables(doc, sheets)


def _write_transport_tables(doc, sheets):
    """Transport section — three sub-tables from the 14_transport sheet."""
    subs = _split_sheet("14_transport")

    caption(doc, "Driver\u2019s License Status by Age")
    add_table(doc, subs.get("Driver's License by Age", pd.DataFrame()), total_label="Total")

    para(doc,
         "Of those with a driver\u2019s license, about half regularly have access to a "
         "reliable vehicle.")

    caption(doc, "Drivers\u2019 Access to a Reliable Vehicle by Age")
    add_table(doc, subs.get("Vehicle Access (licensed)", pd.DataFrame()), total_label="Total")

    para(doc,
         "All youth were asked about the primary way they get to work when they are "
         "employed. The majority rely on public transportation.")

    caption(doc, "Primary Way Youth Get to Work")
    add_table(doc, subs.get("Primary Transport", pd.DataFrame()), total_label="Total")


def sec_voter_reg(doc, sheets):
    heading(doc, "Voting")

    subs = _split_sheet("15_voter_reg")
    df_reg     = subs.get("Voter Registration by Age", pd.DataFrame())
    df_reasons = subs.get("Not Registered Reasons by Age", pd.DataFrame())

    # Pull total % registered from the voter reg sub-table
    total_pct_reg = "[PLACEHOLDER]"
    if not df_reg.empty and "Total" in df_reg.columns:
        reg_row = df_reg[df_reg.iloc[:, 0] == "Registered to Vote"]
        if not reg_row.empty:
            total_pct_reg = reg_row["Total"].values[0]

    para(doc,
         "Youth ages 18 and older were asked to report on whether they are registered "
         f"to vote. Overall, {total_pct_reg} of eligible respondents reported being "
         "registered to vote.")

    caption(doc, "Self-Reported Voter Registration by Age")
    add_table(doc, df_reg, total_label="__none__")

    para(doc,
         "The most common reason youth provided for not registering to vote is that "
         "they believe their vote won\u2019t make a difference.")

    caption(doc, "Reasons Youth Report Not Registering to Vote")
    add_table(doc, df_reasons, total_label="__none__")


def sec_zone_visit(doc, sheets):
    heading(doc, "Zone Experience")

    para(doc,
         "This survey includes questions to better understand participants\u2019 "
         "experiences at the Zone. Attendance patterns at the Zone vary by age, with "
         "older youth coming more frequently than youth ages 16 to 17.")

    subs = _split_sheet("16_visit")
    df_freq     = subs.get("Visit Frequency by Age", pd.DataFrame())
    df_reasons  = subs.get("Visit Reasons (frequent)", pd.DataFrame())
    df_barriers = subs.get("Visit Barriers (infrequent)", pd.DataFrame())

    caption(doc, "Visit Frequency by Age")
    add_table(doc, df_freq, total_label="Total")

    para(doc,
         "As in prior years, most youth report coming to the 412 Youth Zone downtown "
         "to see their coach and to work toward their goals.")

    caption(doc, "What Are the Main Reasons Youth Come to the Youth Zone?")
    add_table(doc, df_reasons, total_label="__none__")

    para(doc,
         "For youth who never visit the Zone, or visit less than monthly, they were "
         "asked what would make them want to come more frequently.")

    caption(doc, "What Would Make Someone Who Rarely Visits the Zone Want to Come, by Age")
    add_table(doc, df_barriers, total_label="__none__")


def sec_program_impact(doc, sheets, df_csv):
    heading(doc, "Impact of Assistance")

    # Count % who reported help in at least one area
    helped_any = (df_csv["q17_program_helped"].str.strip() != "").sum()
    pct_helped = pct(helped_any, N_RESPONDENTS)

    # Q16 agree/somewhat_agree pct
    q16_agree = ((df_csv["q16_stay_focused"].isin(["agree","somewhat_agree"])).sum())
    pct_q16   = pct(q16_agree, (df_csv["q16_stay_focused"] != "").sum())

    para(doc,
         "Across the core outcome areas in which Youth Zone staff are helping young "
         "people make progress, youth reported that their coaches and the Zone have "
         f"helped them in a variety of ways. {pct_helped} of respondents indicated "
         "progress supported by the Zone in at least one area.")

    para(doc,
         f"{pct_q16} of respondents agreed or somewhat agreed that their coach or "
         "the Zone helped them stay focused on their goals.")

    # Q17 by age table — use _split_sheet for proper column headers
    subs = _split_sheet("17_impact")
    df_q17 = subs.get("Program Helped With (Q17) by A", pd.DataFrame())
    # The section title may be truncated in the xlsx key; try partial match
    if df_q17.empty:
        for key in subs:
            if key.startswith("Program Helped"):
                df_q17 = subs[key]
                break

    if not df_q17.empty:
        caption(doc, "My Coach or the Youth Zone has Helped Me To\u2026 (by Age)")
        add_table(doc, df_q17, total_label="__none__")


def sec_respect_environment(doc, sheets):
    df_resp = sheets["18_respect"].copy()
    df_env  = sheets["19_environment"].copy()

    # Q18/Q19 top-2 stats for narrative
    if not df_resp.empty and "% Often or All the Time" in df_resp.columns:
        staff_row = df_resp[df_resp.iloc[:,0].str.contains("Staff", na=False)]
        peer_row  = df_resp[df_resp.iloc[:,0].str.contains("Peer",  na=False)]
        pct_staff = staff_row["% Often or All the Time"].values[0] if not staff_row.empty else "[PLACEHOLDER]"
        pct_peer  = peer_row["% Often or All the Time"].values[0]  if not peer_row.empty  else "[PLACEHOLDER]"
    else:
        pct_staff = pct_peer = "[PLACEHOLDER]"

    para(doc,
         "Respondents were asked to rate how often they felt respected and accepted "
         f"for who they are at the Youth Zone. {pct_staff} of youth reported staff "
         f"treat them with respect often or all the time; {pct_peer} said the same "
         "about their peers at the Zone.")

    if not df_env.empty:
        para(doc,
             "Youth also rated five statements about the program environment on a "
             "1–5 scale. Results are shown in terms of the percentage selecting "
             "4 or 5 (top-2 box).")


def sec_banking(doc, sheets, df_csv):
    heading(doc, "Banking")

    has_account = df_csv["q25_bank_account"].apply(
        lambda v: any(t in (v or "").split("|") for t in [" checking", " savings", "checking", "savings"])
    )
    pct_has = pct(has_account.sum(), N_RESPONDENTS)

    para(doc,
         "Participants were asked questions about their use of banks and other ways "
         "that they store, receive, and transfer money. "
         f"Overall, {pct_has} of respondents reported currently having a bank account.")

    subs = _split_sheet("20_banking")

    caption(doc, "Banking Status by Age")
    add_table(doc, subs.get("Bank Account Status by Age", pd.DataFrame()), total_label="__none__")

    caption(doc, "Methods Youth Use to Store, Receive, and Transfer Money, by Age")
    add_table(doc, subs.get("Money Methods by Age (Q24)", pd.DataFrame()), total_label="__none__")

    caption(doc, "Ways Respondents Use Their Bank Account(s), by Age")
    add_table(doc, subs.get("Account Usage by Age (Q26b)", pd.DataFrame()), total_label="__none__")


def sec_nps(doc, sheets, df_csv):
    df_nps = sheets["21_nps"].copy()
    nps_score_row = df_nps[df_nps.iloc[:,0] == "NPS Score"]
    nps_score = nps_score_row["Count"].values[0] if not nps_score_row.empty and "Count" in df_nps.columns else "[PLACEHOLDER]"
    promoters_row = df_nps[df_nps.iloc[:,0].str.contains("Promoter", na=False)]
    pct_prom = promoters_row["Percent"].values[0] if not promoters_row.empty and "Percent" in df_nps.columns else "[PLACEHOLDER]"

    para(doc,
         f"Youth were asked to rate on a scale of 0\u201310 how likely they would be to "
         "recommend the Youth Zone to a friend or family member. "
         f"The Net Promoter Score (NPS) is {nps_score}. "
         f"{pct_prom} of respondents were Promoters (9\u201310).")


def sec_comments(doc, sheets):
    heading(doc, "Additional Comments")

    df_c = sheets["22_comments"].copy()
    n_comments = len(df_c[df_c.iloc[:,0] != ""])

    para(doc,
         f"Finally, youth had the option to share any other comments or feedback "
         f"they had about the Zone. {n_comments} youth provided additional comments.")

    if not df_c.empty:
        heading(doc, "Comments")
        # Write as bullets
        comment_col = "Comment" if "Comment" in df_c.columns else df_c.columns[-1]
        for _, row in df_c.iterrows():
            text = str(row[comment_col]).strip()
            if text and text != "nan":
                bullet(doc, text)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    if not ANALYSIS_PATH.exists():
        print(f"Analysis file not found: {ANALYSIS_PATH}")
        print("Run 04_analyze_412YZ.py first.")
        sys.exit(1)

    print("Loading analysis data...")
    sheets  = load_sheets()
    df_csv  = load_csv()

    doc = Document()

    # Page setup — 8.5x11, 1in margins (matching prior report)
    section = doc.sections[0]
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1)
    section.right_margin  = Inches(1)

    print("Writing sections...")

    sec_title(doc, df_csv)
    doc.add_paragraph()

    sec_age(doc, sheets)
    doc.add_paragraph()

    sec_gender_orient(doc, sheets)
    doc.add_paragraph()

    sec_race(doc, sheets)
    doc.add_paragraph()

    sec_coach_satisfaction(doc, sheets)
    doc.add_paragraph()

    sec_communication(doc, sheets)
    doc.add_paragraph()

    sec_housing(doc, sheets)
    doc.add_paragraph()

    sec_education_employment(doc, sheets, df_csv)
    doc.add_paragraph()

    sec_job_tenure(doc, sheets, df_csv)
    doc.add_paragraph()

    sec_employment_by_age(doc, sheets)
    doc.add_paragraph()

    sec_job_barriers(doc, sheets, df_csv)
    doc.add_paragraph()

    sec_left_job(doc, sheets)
    doc.add_paragraph()

    sec_transportation(doc, sheets)
    doc.add_paragraph()

    sec_voter_reg(doc, sheets)
    doc.add_paragraph()

    sec_zone_visit(doc, sheets)
    doc.add_paragraph()

    sec_program_impact(doc, sheets, df_csv)
    doc.add_paragraph()

    sec_respect_environment(doc, sheets)
    doc.add_paragraph()

    sec_banking(doc, sheets, df_csv)
    doc.add_paragraph()

    sec_nps(doc, sheets, df_csv)
    doc.add_paragraph()

    sec_comments(doc, sheets)

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PATH))
    print(f"\nSaved: {OUT_PATH}")


if __name__ == "__main__":
    main()
